import streamlit as st
import fitz
from docx import Document
from docx.shared import Pt
import io, re, unicodedata, os, subprocess
from datetime import datetime
from docx2pdf import convert

# --- CONFIGURATION ---
st.set_page_config(page_title="Coco Advisory - Master Portal", layout="centered")
st.title("📂 M&A Process Automation v48")

OUT_DIR = os.path.expanduser("~/Desktop/Process auto/Output/")
if not os.path.exists(OUT_DIR): os.makedirs(OUT_DIR)

def get_d():
    js = ["lundi","mardi","mercredi","jeudi","vendredi","samedi","dimanche"]
    ms = ["janvier","février","mars","avril","mai","juin","juillet","août","septembre","octobre","novembre","décembre"]
    n = datetime.now()
    return f"{js[n.weekday()]} {n.day} {ms[n.month-1]} {n.year}"

def clean_txt(t):
    t = t.lower().replace('\n', ' ')
    t = re.sub(r'\d+/\d+', '', t) 
    t = re.sub(r'\([a-z0-9]+\)', '', t) 
    t = t.replace("$", "")
    t = "".join(c for c in unicodedata.normalize('NFD', t) if unicodedata.category(c) != 'Mn')
    return "".join(re.findall(r'[a-z0-9]+', t))

def find_diff(t1, t2):
    m_len = min(len(t1), len(t2))
    for i in range(m_len):
        if t1[i] != t2[i]:
            s, e = max(0, i - 15), min(m_len, i + 35)
            return t1[s:e], t2[s:e]
    return t1[-30:], t2[-30:]

# --- 1. VALIDATION DU NDA ---
st.subheader("1. Validation du NDA")
u_nda_inv = st.file_uploader("NDA signé par l'investisseur (PDF)", type="pdf")
u_nda_tmp = st.file_uploader("Modèle NDA (Word)", type="docx")

if u_nda_inv and u_nda_tmp:
    if st.button("Vérifier la conformité du NDA"):
        with st.spinner("Analyse en cours..."):
            # Lecture
            d_inv = fitz.open(stream=u_nda_inv.read(), filetype="pdf")
            raw_inv = "".join([pg.get_text("text") for pg in d_inv]); d_inv.close()
            doc_nda_tmp = Document(u_nda_tmp)
            raw_tmp = "\n".join([p.text for p in doc_nda_tmp.paragraphs])
            
            s_mark, e_mark = "messieurs", "sentiments les meilleurs"
            t_l, i_l = raw_tmp.lower(), raw_inv.lower()
            
            # Recherche des positions
            idx_s_t, idx_e_t = t_l.find(s_mark), t_l.find(e_mark)
            idx_s_i, idx_e_i = i_l.find(s_mark), i_l.find(e_mark)
            
            status, diff_msg = "CONFORME", ""
            
            # SÉCURITÉ : Les deux fichiers DOIVENT avoir les marqueurs 
            if any(x == -1 for x in [idx_s_t, idx_e_t, idx_s_i, idx_e_i]):
                status = "ERREUR"
                diff_msg = "Fichiers incompatibles : les marqueurs ('Messieurs' ou 'Sentiments') sont absents."
            else:
                sub_t = clean_txt(raw_tmp[idx_s_t:idx_e_t+len(e_mark)])
                sub_i = clean_txt(raw_inv[idx_s_i:idx_e_i+len(e_mark)])
                
                # SÉCURITÉ : Le texte extrait ne doit pas être dérisoire (ex: moins de 100 caractères)
                if len(sub_t) < 100 or len(sub_i) < 100:
                    status = "ERREUR"
                    diff_msg = "Le texte extrait est trop court. Vérifiez que vous avez chargé les bons fichiers."
                elif sub_t != sub_i:
                    status = "NON CONFORME"
                    f_t, f_i = find_diff(sub_t, sub_i)
                    diff_msg = f"Attendu : ...{f_t}...\nReçu : ...{f_i}..."
            
            # Extraction Société/Signataire
            ls = [l.strip() for l in raw_inv.split('\n') if l.strip()]
            info = {'n': "Inconnu", 's': "Inconnue"}
            for i, l in enumerate(ls):
                if "nom du signataire" in l.lower(): info['n'] = ls[i+1] if i+1 < len(ls) else "Inconnu"
                if "pour le compte de" in l.lower(): info['s'] = ls[i+1] if i+1 < len(ls) else "Inconnue"
            
            st.session_state['nda_res'] = {'status': status, 'diff': diff_msg, 'info': info}

# --- ÉTAPE 2 : SUITE DU PROCESSUS ---
if 'nda_res' in st.session_state:
    res = st.session_state['nda_res']
    if res['status'] == "ERREUR":
        st.error(f"❌ {res['diff']}")
    elif res['status'] == "NON CONFORME":
        st.error(f"🚨 NDA MODIFIÉ : {res['diff']}")
    else:
        st.success(f"✅ NDA CONFORME pour {res['info']['s']}")
        st.markdown("---")
        st.subheader("2. Phase I : Lettre de Process & IM")
        
        u_lp_tmp = st.file_uploader("Modèle Lettre de Process (Word)", type="docx")
        u_im = st.file_uploader("IM filigrané (PDF)", type="pdf")
        
        if u_lp_tmp and u_im:
            doc_lp = Document(u_lp_tmp)
            p_name = "Projet"
            for p in doc_lp.paragraphs:
                if "Projet" in p.text and "Lettre de process" in p.text:
                    m = re.search(r"Projet\s+(.*?)\s*-", p.text); p_name = m.group(1).strip() if m else "Projet"
            
            contact = {"m": "", "p": "Prénom"}
            for t in doc_lp.tables:
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if any(x in cell.text.lower() for x in ["associate", "analyst"]):
                            if r_idx > 0: contact["p"] = t.cell(r_idx-1, c_idx).text.split()[0].capitalize()
                            for o in [1, 2, 3]:
                                if r_idx + o < len(t.rows):
                                    em = t.cell(r_idx+o, c_idx).text.strip()
                                    if "@" in em: contact["m"] = em; break
                            break

            st.write(f"📍 Projet : **{p_name}** | Analyste : **{contact['p']}**")
            adr = st.text_input("Adresse", value="189 rue de Lalande")
            cp = st.text_input("CP / Ville", value="75016 Paris")
            gen = st.selectbox("Genre", ["M", "F"], index=1)
            
            if st.button("Générer le Pack Phase I"):
                p_im = os.path.join(OUT_DIR, u_im.name)
                with open(p_im, "wb") as f: f.write(u_im.getbuffer())
                
                ch = "Cher Monsieur" if gen == "M" else "Chère Madame"
                rems = {"[Destinataire]": res['info']['s'], "[Adresse]": adr, "[Code postale] [Ville]": cp, "[XXX]": res['info']['n'], "[date du jour]": get_d(), "[Cher/Chère] [Monsieur/Madame]": ch}
                
                for p in list(doc_lp.paragraphs) + [p for t in doc_lp.tables for r_w in t.rows for c in r_w.cells for p in c.paragraphs]:
                    for k, v in rems.items():
                        if k in p.text: p.text = p.text.replace(k, v)
                    for run in p.runs:
                        run.font.name = 'Calibri Light'; run.font.size = Pt(10); run.font.highlight_color = None
                
                cl = res['info']['s'].replace("/", "-")
                p_w, p_p = os.path.join(OUT_DIR, f"LP-{cl}.docx"), os.path.join(OUT_DIR, f"LP-{cl}.pdf")
                doc_lp.save(p_w); convert(p_w, p_p)
                st.session_state['final'] = {'paths': (p_w, p_p, p_im), 'proj': p_name, 'contact': contact, 'soc': res['info']['s']}
                st.success("Pack générée (Calibri 10).")

        if 'final' in st.session_state:
            if st.button("📧 Envoyer le mail automatique"):
                f = st.session_state['final']
                suj = f"Projet {f['proj']} - LP - {f['soc']}"
                txt = f"Bonjour {f['contact']['p']},\\n\\nLe NDA est bien conforme. Tu trouveras en pièce jointe les fichiers de LP (Word et PDF) ainsi que l'IM filigrané.\\n\\nBien à toi,\\nCorentin GUE"
                sc = f'tell application "Mail"\nactivate\nset msg to make new outgoing message with properties {{subject:"{suj}", content:"{txt}", visible:true}}\ntell msg\nmake new recipient at end of to recipients with properties {{address:"{f["contact"]["m"]}"}}\nmake new attachment with properties {{file name:POSIX file "{f["paths"][0]}"}} at after the last paragraph\nmake new attachment with properties {{file name:POSIX file "{f["paths"][1]}"}} at after the last paragraph\nmake new attachment with properties {{file name:POSIX file "{f["paths"][2]}"}} at after the last paragraph\nend tell\nend tell'
                subprocess.run(['osascript', '-e', sc])
                st.balloons()